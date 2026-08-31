import os
import json
import re
import uuid
import io
from datetime import datetime
from flask import Blueprint, request, jsonify, render_template, current_app, send_file, render_template_string, make_response, Response, stream_with_context
from pydantic import ValidationError
from sqlalchemy.sql.expression import func
from sqlalchemy.orm import joinedload
from docx import Document
from docx.shared import Inches
from weasyprint import HTML
import pathlib
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from .models import Question, Paper, PaperQuestion, Visitor, Board, Class, Subject, Chapter
from .ai_providers import (generate_with_fallback, stream_with_fallback, ProviderError,
                           available_providers, provider_status)
from .schemas import GenerateRequest, AIQuestion
from .prompts import build_topic_prompt, build_class_prompt, build_language_instruction
from . import db, limiter


main = Blueprint("main", __name__)


def _debug_enabled() -> bool:
    return os.getenv("FLASK_DEBUG", "").lower() in ("1", "true", "yes")


@main.route("/db-test")
def db_test():
    if not _debug_enabled():
        return jsonify({"error": "not found"}), 404
    try:
        from sqlalchemy import text
        db.session.execute(text("SELECT 1"))
        return "Database connected successfully."
    except Exception as e:
        return f"DB error: {e}", 500


@main.route("/debug-questions")
def debug_questions():
    if not _debug_enabled():
        return jsonify({"error": "not found"}), 404
    return {"count": Question.query.count()}





# In app/routes.py (After imports, before get_or_create_visitor)

# Helper function to clean and render simple math symbols for WeasyPrint/ReportLab
def render_simple_math(text):
    if not text:
        return ""
    
    # 1. Remove surrounding $ signs (which the AI uses for LaTeX)
    text = re.sub(r'^\s*\$', '', text)
    text = re.sub(r'\$\s*$', '', text)
    
    # 2. Replace common LaTeX commands with Unicode/HTML symbols
    text = text.replace(r'\times', '×')      # Multiplication symbol
    text = text.replace(r'^\circ', '°')       # Degree symbol (used in 40^\circ)
    text = text.replace(r'\circ', '°')        # Degree symbol
    text = text.replace(r'\le', '≤')
    text = text.replace(r'\ge', '≥')
    text = text.replace(r'\ne', '≠')
    
    # 3. Handle simple superscripts (e.g., 3^2 -> 3<sup>2</sup>)
    # This is a simplification; WeasyPrint supports <sup> tags.
    text = re.sub(r'(\d+)\^(\d+)', r'\1<sup>\2</sup>', text)
    
    # 4. Handle simple fractions: replace \frac{1}{3} with (1/3) or <sup>1</sup>&frasl;<sub>3</sub>
    # Using simple HTML for better readability in PDF
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'<sup>\1</sup>&frasl;<sub>\2</sub>', text)
    
    return text







def get_or_create_visitor():
    """Get existing visitor from cookie or create new visitor"""
    visitor_id = request.cookies.get('visitor_id')
    
    if visitor_id:
        # Check if visitor exists in database
        visitor = Visitor.query.filter_by(visitor_id=visitor_id).first()
        if visitor:
            # Update last visit and increment visit count
            visitor.visit_count += 1
            db.session.commit()
            return visitor
    
    # Create new visitor
    visitor_id = str(uuid.uuid4())
    visitor = Visitor()
    visitor.visitor_id = visitor_id
    db.session.add(visitor)
    db.session.commit()
    
    return visitor




# In app/routes.py, replace your get_academic_data function with this

# Subjects withheld from the picker. The paper generator is English-only, so the
# Hindi-language subject is not offered; its rows remain in the database.
_HIDDEN_SUBJECTS = {"hindi"}


def _is_hidden_subject(name_en: str) -> bool:
    return (name_en or "").strip().lower() in _HIDDEN_SUBJECTS


@main.route("/api/academic-data")
def get_academic_data():
    """
    Queries the database and formats the data to perfectly match
    the original hardcoded JavaScript object structure.
    """
    try:
        # THIS IS THE CHANGED LINE: Use joinedload for an efficient, single query
        all_boards = Board.query.options(
            joinedload(Board.classes).joinedload(Class.subjects).joinedload(Subject.chapters)
        ).all()

        # Initialize the two main dictionaries that our frontend expects
        subjects_data = {}
        chapters_data = {}

        for board in all_boards:
            subjects_data[board.name] = {}
            chapters_data[board.name] = {}
            sorted_classes = sorted(board.classes, key=lambda c: c.class_number)

            for class_ in sorted_classes:
                class_num_str = str(class_.class_number)
                subjects_data[board.name][class_num_str] = []
                chapters_data[board.name][class_num_str] = {}

                for subject in class_.subjects:
                    # Hindi is hidden from the subject list: the app generates
                    # English-only papers, and the Hindi subject carries no
                    # questions in the bank. The rows stay in the database, so
                    # deleting this check brings the subject straight back.
                    if _is_hidden_subject(subject.name_en):
                        continue
                    subjects_data[board.name][class_num_str].append({
                        "en": subject.name_en,
                        "hi": subject.name_hi
                    })
                    chapters_list = []
                    for chapter in subject.chapters:
                        chapters_list.append({
                            "en": chapter.title_en,
                            "hi": chapter.title_hi
                        })
                    chapters_data[board.name][class_num_str][subject.name_en] = chapters_list

        return jsonify({
            "subjects": subjects_data,
            "chapters": chapters_data
        })

    except Exception as e:
        current_app.logger.error(f"Failed to fetch and format academic data: {e}")
        return jsonify({"error": "Could not fetch academic data."}), 500





@main.route("/api/ai-status")
def ai_status():
    """Health of the AI provider, polled by the page after a failed generation."""
    return jsonify(provider_status())


@main.route("/")
def index():
    response = make_response(render_template("index.html", ai_status=provider_status()))
    
    # Check if visitor cookie exists
    visitor_id = request.cookies.get('visitor_id')
    if not visitor_id:
        # Create new visitor and set cookie
        visitor = get_or_create_visitor()
        response.set_cookie('visitor_id', visitor.visitor_id, max_age=365*24*60*60)  # 1 year
    
    return response

# Helper function to normalize question types
# Canonical key -> section name. Keys are lowercased with every non-alphanumeric
# character removed, so "Multiple Choice", "multiple-choice" and "multiplechoice"
# all collapse to the same entry. LLMs are inconsistent about punctuation and
# casing here, and an unmatched type never lands in a PDF section, so the
# question would silently vanish from the paper.
_QTYPE_CANON = {
    "mcq": "MCQ",
    "multiplechoice": "MCQ",
    "multiplechoicequestion": "MCQ",
    "fill": "Fill in the Blanks",
    "fillintheblank": "Fill in the Blanks",
    "fillintheblanks": "Fill in the Blanks",
    "short": "Short Answer",
    "shortanswer": "Short Answer",
    "long": "Long Answer",
    "longanswer": "Long Answer",
    "match": "Matching",
    "matching": "Matching",
    "matchthefollowing": "Matching",
    "case": "Case Study",
    "casestudy": "Case Study",
}


def _normalize_qtype(label):
    label = (label or "").strip()
    key = re.sub(r"[^a-z0-9]", "", label.lower())
    return _QTYPE_CANON.get(key, label)

def _build_qdist_string(qdist: dict) -> str:
    """Verbose, count-aware breakdown so the LLM is less likely to under-deliver."""
    parts = []
    grand_total_q = 0
    grand_total_marks = 0
    for qtype, info in qdist.items():
        if not isinstance(info, dict):
            continue
        count = int(info.get("count", 0) or 0)
        marks = int(info.get("marks", 0) or 0)
        if count > 0:
            parts.append(f"- EXACTLY {count} {qtype} question(s), each worth {marks} marks "
                         f"(subtotal: {count * marks} marks)")
            grand_total_q += count
            grand_total_marks += count * marks
    if grand_total_q:
        parts.append(f"\nGRAND TOTAL: {grand_total_q} questions, {grand_total_marks} marks. "
                     f"You MUST return exactly {grand_total_q} questions — no fewer.")
    return "\n".join(parts)


def _parse_ai_json(raw_text: str) -> list:
    """Parse the LLM's response into a list of question dicts.

    Handles raw JSON, fenced ```json blocks, and best-effort recovery.
    """
    if not raw_text:
        raise ValueError("AI returned empty response.")
    raw_text = raw_text.strip()
    try:
        data = json.loads(raw_text)
    except json.JSONDecodeError:
        m = re.search(r"```(?:json)?\s*(.*?)```", raw_text, re.DOTALL)
        if m:
            data = json.loads(m.group(1))
        else:
            m2 = re.search(r"\[\s*\{.*\}\s*\]", raw_text, re.DOTALL)
            if not m2:
                raise ValueError("AI returned invalid JSON.")
            data = json.loads(m2.group(0))
    if isinstance(data, dict):
        for key in ("questions", "data", "items"):
            if key in data and isinstance(data[key], list):
                data = data[key]
                break
    if not isinstance(data, list):
        raise ValueError("AI response was not a JSON array of questions.")
    return data


def _validate_questions(raw_list: list, paper_language: str) -> list[dict]:
    """Validate each AI question through Pydantic; drop bad ones."""
    out = []
    for raw in raw_list:
        if not isinstance(raw, dict):
            continue
        try:
            q = AIQuestion(**raw).model_dump()
        except ValidationError as ve:
            current_app.logger.warning("Dropping malformed AI question: %s", ve)
            continue
        q["question_type"] = _normalize_qtype(q.get("type"))
        q["question_text"] = q.get("question", "")
        q["source"] = "AI"
        q["language"] = paper_language
        if q["question_type"] != "MCQ":
            q["options"] = None
        # Models routinely mislabel these two structured types, or put the
        # case-study passage in "question" and leave "passage" empty. Trust the
        # content over the label, then strip whatever does not belong to the
        # resulting type — otherwise a case-study passage renders underneath a
        # Match the Following question.
        has_pairs = bool(q.get("pairs"))
        has_case = bool(q.get("passage") or q.get("sub_questions"))
        if q["question_type"] == "Case Study" and has_pairs and not has_case:
            q["question_type"] = "Matching"
        elif q["question_type"] == "Matching" and has_case and not has_pairs:
            q["question_type"] = "Case Study"

        if q["question_type"] != "Matching":
            q["pairs"] = None
        if q["question_type"] != "Case Study":
            q["passage"] = None
            q["sub_questions"] = None
        # Deliberately NOT inventing pairs or a passage when they are missing:
        # the question text usually already carries the content, and fabricated
        # filler in a real exam paper is worse than a plainer question.
        out.append(q)
    return out


# --- Smart backfill: AI retry -> loosened DB lookup -> typed placeholders ---


def _compute_shortfall(questions: list, qdist: dict) -> dict:
    """Returns {normalized_type: {needed, have, missing, marks, qtype_frontend}} for any type still short."""
    sf = {}
    for qtype_frontend, info in qdist.items():
        normalized = _normalize_qtype(qtype_frontend)
        needed = int(info.get("count", 0))
        marks = int(info.get("marks", 1))
        have = sum(1 for q in questions if q.get("question_type") == normalized)
        if have < needed:
            sf[normalized] = {
                "needed": needed, "have": have,
                "missing": needed - have, "marks": marks,
                "qtype_frontend": qtype_frontend,
            }
    return sf


def _retry_ai_for_missing(shortfall: dict, ctx: dict) -> list:
    """Targeted AI retry asking only for the missing items."""
    if not shortfall:
        return []
    parts = [
        f"- exactly {sf['missing']} {sf['qtype_frontend']} question(s), each worth {sf['marks']} marks"
        for sf in shortfall.values()
    ]
    scope_bits = []
    if ctx.get("topic"):
        scope_bits.append(f'focused on topic "{ctx["topic"]}"')
    elif ctx.get("chapters"):
        scope_bits.append(f"covering chapters: {', '.join(ctx['chapters'])}")
    scope = " ".join(scope_bits) or "from the general syllabus"
    retry_prompt = (
        f"You previously generated questions for a Class {ctx['class_']} {ctx['subject']} "
        f"({ctx['board']}) paper. I still need MORE questions. Generate ONLY these MISSING questions, "
        f"{scope}.\n\nMissing:\n" + "\n".join(parts) +
        f"\n\nLanguage: {ctx['paper_language']}.\n"
        "Return ONLY a JSON array. Each object: type, question, options (MCQ only — exactly 4), "
        "marks (integer), difficulty (Easy/Medium/Hard), answer, explanation.\n"
        'For "Matching" you MUST also include "pairs": exactly 4 [term, definition] entries.\n'
        'For "Case Study" you MUST also include "passage" (a 40-100 word stimulus the '
        'student reads) and "sub_questions" (2-4 questions about it). Put the passage in '
        '"passage", NOT in "question" — "question" is only the instruction line.'
    )
    try:
        text, used = generate_with_fallback(retry_prompt)
        current_app.logger.info("AI retry-for-missing succeeded via %s", used)
        raw = _parse_ai_json(text)
        return _validate_questions(raw, ctx["paper_language"])
    except Exception as e:
        current_app.logger.warning("AI retry-for-missing failed: %s", e)
        return []


def _placeholder_pairs(paper_language: str, scope: str = "") -> list:
    """Four Column A / Column B rows, used when a Matching question arrives empty."""
    where = f" ({scope})" if scope else ""
    if (paper_language or "").lower() == "hindi":
        return [
            [f"पद 1{where}", "इस पद की परिभाषा"],
            ["पद 2", "इस पद की परिभाषा"],
            ["पद 3", "इस पद की परिभाषा"],
            ["पद 4", "इस पद की परिभाषा"],
        ]
    return [
        [f"Term 1{where}", "Definition of the first term"],
        ["Term 2", "Definition of the second term"],
        ["Term 3", "Definition of the third term"],
        ["Term 4", "Definition of the fourth term"],
    ]


def _placeholder_case_study(paper_language: str, scope: str = "") -> tuple:
    """A stimulus passage plus sub-questions, for an empty Case Study question."""
    topic = scope or ("इस अध्याय" if (paper_language or "").lower() == "hindi" else "this chapter")
    if (paper_language or "").lower() == "hindi":
        return (
            f"[गद्यांश शिक्षक द्वारा जोड़ा जाना है — {topic} पर आधारित केस स्टडी।]",
            [
                "गद्यांश में दी गई मुख्य अवधारणा क्या है?",
                "इस अवधारणा का एक व्यावहारिक उपयोग बताइए।",
                "अपने उत्तर की पुष्टि कीजिए।",
            ],
        )
    return (
        f"[Passage to be added by the teacher — case study on {topic}.]",
        [
            "Identify the main concept described in the passage.",
            "Give one practical application of this concept.",
            "Justify your answer with a short explanation.",
        ],
    )


def _smart_placeholder(qtype_frontend: str, marks: int, subject: str, class_: str,
                       paper_language: str, scope_hint: str = "") -> dict:
    """Typed placeholder respecting marks. Scope hint = topic or 'general syllabus'."""
    qtype_norm = _normalize_qtype(qtype_frontend)
    scope = scope_hint or f"Class {class_} {subject}"
    base = {
        "type": qtype_frontend,
        "question_type": qtype_norm,
        "marks": marks,
        "difficulty": "Medium",
        "source": "Fallback",
        "language": paper_language,
        "options": None,
        "answer": "See your textbook for the model answer.",
        "explanation": "Refer to the relevant chapter in your textbook.",
    }
    if qtype_norm == "MCQ":
        qtext = f"Which of the following best describes a key concept from {scope}?"
        return {**base, "question": qtext, "question_text": qtext,
                "options": ["Option A", "Option B", "Option C", "Option D"], "answer": "A"}
    if qtype_norm == "Fill in the Blanks":
        qtext = f"In {scope}, the concept of ________ plays an important role."
        return {**base, "question": qtext, "question_text": qtext}
    if qtype_norm == "Long Answer":
        qtext = f"Explain in detail an important topic from {scope}, including examples and applications."
        return {**base, "question": qtext, "question_text": qtext}
    if qtype_norm == "Matching":
        qtext = (
            "स्तंभ A की प्रविष्टियों का स्तंभ B से मिलान कीजिए।"
            if paper_language == "hindi"
            else "Match the entries in Column A with the correct entries in Column B."
        )
        return {**base, "question": qtext, "question_text": qtext,
                "pairs": _placeholder_pairs(paper_language, scope)}
    if qtype_norm == "Case Study":
        passage, subs = _placeholder_case_study(paper_language, scope)
        qtext = (
            "निम्नलिखित गद्यांश पढ़िए और उसके नीचे दिए गए प्रश्नों के उत्तर दीजिए।"
            if paper_language == "hindi"
            else "Read the passage below and answer the questions that follow."
        )
        return {**base, "question": qtext, "question_text": qtext,
                "passage": passage, "sub_questions": subs}
    qtext = f"Briefly describe an important concept from {scope}."
    return {**base, "question": qtext, "question_text": qtext}


def _backfill_distribution(*, questions: list, qdist: dict, board: str, class_: str,
                            subject: str, chapters: list, topic: str,
                            paper_language: str) -> list:
    """Top up `questions` to meet per-type counts. AI retry -> DB (loosened) -> placeholder."""
    shortfall = _compute_shortfall(questions, qdist)
    if not shortfall:
        return questions

    current_app.logger.info(
        "Distribution shortfall before backfill: %s",
        {k: v["missing"] for k, v in shortfall.items()},
    )

    # 1) Targeted AI retry for the missing items.
    ctx = {"board": board, "class_": class_, "subject": subject,
           "topic": topic, "chapters": chapters, "paper_language": paper_language}
    retry_questions = _retry_ai_for_missing(shortfall, ctx)
    existing_texts = {q.get("question_text", "").strip().lower() for q in questions}
    for q in retry_questions:
        nt = q.get("question_type")
        if nt not in shortfall or shortfall[nt]["missing"] <= 0:
            continue
        text = (q.get("question_text") or "").strip().lower()
        if not text or text in existing_texts:
            continue
        # Force expected marks so totals match what the user asked for.
        q["marks"] = shortfall[nt]["marks"]
        questions.append(q)
        existing_texts.add(text)
        shortfall[nt]["missing"] -= 1
    shortfall = {k: v for k, v in shortfall.items() if v["missing"] > 0}
    if not shortfall:
        return questions

    # 2) DB lookup with progressive filter loosening.
    topic_present = bool(topic and topic.strip())
    chapter_id_list = []
    if chapters:
        chapter_objects = Chapter.query.filter(Chapter.title_en.in_(chapters)).all()
        chapter_id_list = [c.chapter_id for c in chapter_objects]

    def _build_base_query(ntype, marks=None, language=None):
        f = {"question_type": ntype}
        if marks is not None:
            f["marks"] = marks
        if language is not None:
            f["language"] = language
        q = Question.query.filter_by(**f)
        if chapter_id_list:
            return q.filter(Question.chapter_id.in_(chapter_id_list))
        if topic_present:
            return q.filter(Question.question_text.contains(topic))
        q = q.join(Chapter).join(Subject)
        if subject:
            q = q.filter(Subject.name_en == subject)
        if class_:
            q = q.join(Class).filter(Class.class_number == class_)
        return q

    for ntype, sf in list(shortfall.items()):
        if sf["missing"] <= 0:
            continue
        seen_ids = set()
        filter_variants = [
            dict(marks=sf["marks"], language=paper_language),  # exact
            dict(language=paper_language),                     # any marks, same lang
            dict(marks=sf["marks"]),                           # exact marks, any lang
            dict(),                                            # any marks, any lang
        ]
        for variant in filter_variants:
            if sf["missing"] <= 0:
                break
            try:
                rows = (
                    _build_base_query(ntype, **variant)
                    .order_by(func.rand())
                    .limit(sf["missing"] * 3)
                    .all()
                )
            except Exception as e:
                current_app.logger.warning("DB backfill query failed (%s, %s): %s", ntype, variant, e)
                rows = []
            for r in rows:
                if sf["missing"] <= 0:
                    break
                if r.id in seen_ids:
                    continue
                seen_ids.add(r.id)
                text = (r.question_text or "").strip().lower()
                if not text or text in existing_texts:
                    continue
                q_dict = r.as_dict()
                q_dict["source"] = "Database"
                q_dict["language"] = r.language or paper_language
                q_dict["question_type"] = ntype
                q_dict["marks"] = sf["marks"]  # honor user's requested marks
                q_dict["question_text"] = r.question_text
                questions.append(q_dict)
                existing_texts.add(text)
                sf["missing"] -= 1
        current_app.logger.info("After DB backfill type=%s missing=%d", ntype, sf["missing"])

    shortfall = {k: v for k, v in shortfall.items() if v["missing"] > 0}
    if not shortfall:
        return questions

    # 3) Typed placeholders for whatever is still missing.
    scope_hint = ""
    if topic_present:
        scope_hint = topic
    elif chapters:
        scope_hint = f"the chapter '{chapters[0]}' of Class {class_} {subject}"
    for ntype, sf in shortfall.items():
        for _ in range(sf["missing"]):
            questions.append(
                _smart_placeholder(sf["qtype_frontend"], sf["marks"], subject,
                                   class_, paper_language, scope_hint)
            )
        current_app.logger.info("Smart placeholders added: type=%s count=%d", ntype, sf["missing"])

    return questions


@main.route("/api/generate", methods=["POST"])
@limiter.limit("10 per hour")
def generate_paper():
    raw_payload = request.get_json(silent=True) or {}
    try:
        req = GenerateRequest(**raw_payload)
    except ValidationError as ve:
        return jsonify({"error": "Invalid request", "details": ve.errors()}), 400

    subject = req.subject
    class_ = req.class_
    board = req.schoolBoard
    school = req.schoolName or ""
    qdist = {k: v.model_dump() for k, v in req.questionDistribution.items()}
    ddist = req.difficultyDistribution
    exam_name = req.examName or ""
    paper_language = req.paperLanguage
    topic = req.topic or ""
    chapters = req.chapters
    questions: list = []

    visitor = get_or_create_visitor()
    topic_present = bool(topic and topic.strip())
    current_app.logger.info(
        "generate_paper subject=%s class=%s board=%s topic_present=%s providers=%s",
        subject, class_, board, topic_present,
        [p.name for p in available_providers()],
    )

    qdist_prompt_str = _build_qdist_string(qdist)
    language_instruction = build_language_instruction(paper_language, subject)

    if topic_present:
        prompt = build_topic_prompt(
            board=board, school=school, class_=class_, subject=subject,
            topic=topic, qdist_str=qdist_prompt_str, ddist=ddist,
            language_instruction=language_instruction, paper_language=paper_language,
        )
    else:
        prompt = build_class_prompt(
            board=board, school=school, class_=class_, subject=subject,
            chapters=chapters, qdist_str=qdist_prompt_str, ddist=ddist,
            language_instruction=language_instruction, paper_language=paper_language,
        )

    try:
        raw_text, provider_used = generate_with_fallback(prompt)
        current_app.logger.info("AI generation succeeded via provider=%s", provider_used)
        ai_questions_raw = _parse_ai_json(raw_text)
        questions = _validate_questions(ai_questions_raw, paper_language)
        
        # Track used question IDs to prevent duplicates
        used_question_hashes = set()
        
        balanced_questions = []
        for qtype_frontend, info in qdist.items():
            count_needed = int(info['count'])
            normalized_type = _normalize_qtype(qtype_frontend)
            # Filter questions of this type
            type_questions = [q for q in questions if q.get("question_type") == normalized_type]
            
            # Add questions to balanced list, avoiding duplicates
            added_count = 0
            for q in type_questions:
                # Check if we've already added enough questions of this type
                if added_count >= count_needed:
                    break
                    
                # Create a simple hash of the question text to detect duplicates
                question_hash = hash(q.get("question", "").strip().lower())
                if question_hash not in used_question_hashes:
                    balanced_questions.append(q)
                    used_question_hashes.add(question_hash)
                    added_count += 1
        
        questions = balanced_questions
        
        # PASTE THIS NEW BLOCK IN ITS PLACE:

        # Find a suitable chapter_id to associate these new questions with.
        first_chapter_id = None
        if not topic_present and chapters:
             # Find the first chapter that matches the user's selection
             first_chapter_object = Chapter.query.join(Subject).join(Class).join(Board).filter(
                Board.name == board,
                Class.class_number == class_,
                Subject.name_en == subject,
                Chapter.title_en == chapters[0]
            ).first()
             if first_chapter_object:
                 first_chapter_id = first_chapter_object.chapter_id

        # This block now correctly saves questions with a chapter_id
        processed_questions = []
        for q in questions: # Note: 'questions' here is the list from the AI
            q_type = _normalize_qtype(q.get("type"))
            q_text = q.get("question", "")
            
            # Only save to DB if we found a chapter to link it to
            if first_chapter_id:
                new_q = Question(
                    chapter_id=first_chapter_id,
                    question_type=q_type,
                    difficulty=q.get("difficulty"),
                    marks=q.get("marks"),
                    question_text=q_text,
                    options=q.get("options") if q_type == "MCQ" else None,
                    answer=q.get("answer", "Not provided"),
                    source="AI",
                    explanation=q.get("explanation", ""),
                    language=paper_language
                )
                db.session.add(new_q)
                db.session.flush() # Use flush to get the ID before committing
                q['id'] = new_q.id # Add the new database ID to the question object

            q['source'] = "AI"
            q['question_type'] = q_type
            q['question_text'] = q_text
            processed_questions.append(q)

        questions = processed_questions
        db.session.commit()

    except Exception as e:
        current_app.logger.error(f"AI generation failed: {e}. Falling back completely to DB.")
        db.session.rollback()
        questions = []

    # --- Smart backfill: AI retry -> DB (loosened) -> typed placeholders ---
    questions = _backfill_distribution(
        questions=questions, qdist=qdist, board=board, class_=class_,
        subject=subject, chapters=chapters, topic=topic, paper_language=paper_language,
    )
    # Persist + render through the shared helper the streaming route also uses.
    # This path used to carry its own copy of the DB write, JSON dump and PDF
    # template, which had drifted: English-only section titles, always-plural
    # "marks", and no Devanagari bold face. One renderer, one look.
    paper_id, urls, summary = _persist_and_render_paper(
        questions=questions,
        board=board,
        class_=class_,
        subject=subject,
        school=school,
        exam_name=exam_name,
        paper_language=paper_language,
        visitor_id=visitor.visitor_id,
        qdist=qdist,
    )

    type_order = ["MCQ", "Multiple Choice", "Fill in the Blanks", "Fill", "Short Answer",
                  "Short", "Long Answer", "Long", "Matching", "Match",
                  "Match the Following", "Case Study", "Case"]

    def _order(q):
        t = (q.get("type") or q.get("question_type") or "").strip()
        try:
            return type_order.index(t)
        except ValueError:
            return len(type_order)

    questions_sorted = sorted(questions, key=_order)

    return jsonify({
        "questions": [{
            "id": q.get("id"), "question_text": q.get("question_text"),
            "marks": q.get("marks"), "difficulty": q.get("difficulty"),
            "type": q.get("type") or q.get("question_type"),
            "source": q.get("source", "Database"),
            "options": q.get("options"), "pairs": q.get("pairs"),
            "passage": q.get("passage"), "sub_questions": q.get("sub_questions"),
        } for q in questions_sorted],
        "summary": summary,
        "pdf_url": urls["pdf_url"],
        "word_url": urls["word_url"],
        "answer_key_url": urls["answer_key_url"],
    })


def _sse(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


@main.route("/api/generate/stream", methods=["POST"])
@limiter.limit("10 per hour")
def generate_paper_stream():
    """Server-Sent Events endpoint: stream questions as the LLM produces them.

    Emits events:
      - phase   { phase: "connecting" | "generating" | "rendering" | "done" }
      - chunk   { text: "...partial llm output..." }
      - question{ index, question, provider }      (each parsed question)
      - error   { message }
      - done    { paper_id, pdf_url, word_url, answer_key_url, summary }
    """
    raw_payload = request.get_json(silent=True) or {}
    try:
        req = GenerateRequest(**raw_payload)
    except ValidationError as ve:
        return jsonify({"error": "Invalid request", "details": ve.errors()}), 400

    if not available_providers():
        return jsonify({"error": "No AI provider configured."}), 503

    visitor = get_or_create_visitor()
    visitor_id = visitor.visitor_id

    subject = req.subject
    class_ = req.class_
    board = req.schoolBoard
    school = req.schoolName or ""
    qdist = {k: v.model_dump() for k, v in req.questionDistribution.items()}
    ddist = req.difficultyDistribution
    exam_name = req.examName or ""
    paper_language = req.paperLanguage
    topic = req.topic or ""
    chapters = req.chapters
    topic_present = bool(topic and topic.strip())

    qdist_prompt_str = _build_qdist_string(qdist)
    language_instruction = build_language_instruction(paper_language, subject)
    if topic_present:
        prompt = build_topic_prompt(
            board=board, school=school, class_=class_, subject=subject,
            topic=topic, qdist_str=qdist_prompt_str, ddist=ddist,
            language_instruction=language_instruction, paper_language=paper_language,
        )
    else:
        prompt = build_class_prompt(
            board=board, school=school, class_=class_, subject=subject,
            chapters=chapters, qdist_str=qdist_prompt_str, ddist=ddist,
            language_instruction=language_instruction, paper_language=paper_language,
        )

    app_obj = current_app._get_current_object()

    def generator():
        with app_obj.app_context():
            try:
                yield _sse("phase", {"phase": "connecting"})
                buffer = ""
                provider_seen = None
                yielded_count = 0
                yield _sse("phase", {"phase": "generating"})
                try:
                    for chunk, provider_name in stream_with_fallback(prompt):
                        provider_seen = provider_name
                        buffer += chunk
                        yield _sse("chunk", {"text": chunk, "provider": provider_name})
                        # Best-effort progressive parse of completed top-level JSON objects.
                        new_count, parsed_qs = _extract_complete_questions(buffer, yielded_count)
                        for q in parsed_qs:
                            yield _sse("question", {"index": yielded_count, "question": q, "provider": provider_name})
                            yielded_count += 1
                        yielded_count = new_count
                except ProviderError as pe:
                    yield _sse("error", {"message": f"All providers failed: {pe}"})
                    return

                yield _sse("phase", {"phase": "rendering"})
                try:
                    ai_questions_raw = _parse_ai_json(buffer)
                    questions = _validate_questions(ai_questions_raw, paper_language)
                except Exception as parse_err:
                    current_app.logger.warning("Streaming parse failed, will rely on backfill: %s", parse_err)
                    questions = []

                # Top up to meet the requested distribution.
                questions = _backfill_distribution(
                    questions=questions, qdist=qdist, board=board, class_=class_,
                    subject=subject, chapters=chapters, topic=topic, paper_language=paper_language,
                )

                if not questions:
                    yield _sse("error", {"message": "Could not generate any questions."})
                    return

                paper_id, urls, summary = _persist_and_render_paper(
                    questions=questions,
                    board=board, class_=class_, subject=subject,
                    school=school, exam_name=exam_name,
                    paper_language=paper_language, visitor_id=visitor_id,
                    qdist=qdist,
                )
                questions_payload = [{
                    "question_text": q.get("question_text") or q.get("question"),
                    "question": q.get("question_text") or q.get("question"),
                    "marks": q.get("marks"),
                    "difficulty": q.get("difficulty"),
                    "type": q.get("type") or q.get("question_type"),
                    "question_type": q.get("question_type"),
                    "options": q.get("options"),
                    "answer": q.get("answer"),
                    # Without these the preview card shows only the instruction
                    # line for Matching and Case Study questions, and the editor
                    # has nothing to edit.
                    "pairs": q.get("pairs"),
                    "passage": q.get("passage"),
                    "sub_questions": q.get("sub_questions"),
                } for q in questions]
                yield _sse("done", {
                    "paper_id": paper_id,
                    "summary": summary,
                    "provider": provider_seen,
                    "questions": questions_payload,
                    **urls,
                })
            except Exception as e:
                current_app.logger.exception("stream generator failed")
                yield _sse("error", {"message": str(e)})

    return Response(
        stream_with_context(generator()),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


_QUESTION_OBJ_RE = re.compile(r"\{[^{}]*?\"question\"\s*:\s*\"[^\"]*\"[^{}]*?\}", re.DOTALL)


def _extract_complete_questions(buffer: str, already_yielded: int):
    """Best-effort extractor for completed top-level JSON objects in a streaming buffer.

    Returns (new_yielded_count, list_of_newly_parsed_dicts).
    Very forgiving — failures just mean we wait for more data.
    """
    matches = _QUESTION_OBJ_RE.findall(buffer)
    out = []
    for raw in matches[already_yielded:]:
        try:
            obj = json.loads(raw)
            if isinstance(obj, dict) and "question" in obj:
                out.append(obj)
        except Exception:
            continue
    return already_yielded + len(out), out


# Section headings, without their letter. The letter is assigned at render time
# over the sections that actually have questions — otherwise a paper with no MCQs
# would open at "Section B".
_SECTION_ORDER = ["Multiple Choice", "Fill in the Blanks", "Short Answer",
                  "Long Answer", "Matching", "Case Study"]

_SECTION_NAMES = {
    "english": {
        "Multiple Choice":    "Multiple Choice Questions",
        "Fill in the Blanks": "Fill in the Blanks",
        "Short Answer":       "Short Answer Questions",
        "Long Answer":        "Long Answer Questions",
        "Matching":           "Match the Following",
        "Case Study":         "Case Study",
    },
    "hindi": {
        "Multiple Choice":    "बहुविकल्पीय प्रश्न",
        "Fill in the Blanks": "रिक्त स्थान भरें",
        "Short Answer":       "लघु उत्तरीय प्रश्न",
        "Long Answer":        "दीर्घ उत्तरीय प्रश्न",
        "Matching":           "मिलान करें",
        "Case Study":         "केस अध्ययन",
    },
}

_SECTION_WORD = {"english": "Section", "hindi": "खंड"}

_PDF_LABELS = {
    "english": {
        "q_prefix":     "Q",
        "marks_one":    "mark",
        "marks_many":   "marks",
        "date":         "Date",
        "class":        "Class",
        "answer_key":   "Answer Key",
        "answer":       "Answer",
        "explanation":  "Explanation",
        "column_a":     "Column A",
        "column_b":     "Column B",
    },
    "hindi": {
        "q_prefix":     "प्र.",
        "marks_one":    "अंक",
        "marks_many":   "अंक",
        "date":         "दिनांक",
        "class":        "कक्षा",
        "answer_key":   "उत्तर कुंजी",
        "answer":       "उत्तर",
        "explanation":  "व्याख्या",
        "column_a":     "स्तंभ A",
        "column_b":     "स्तंभ B",
    },
}


def _section_choice_notes(qdist: dict | None, sections: dict, paper_language: str) -> dict:
    """Per-section "Attempt any N of the M questions" lines.

    Driven by the requested distribution rather than the rendered count, but
    capped at what actually got printed — the backfill can come up short, and
    telling a student to answer 4 of 6 when only 5 exist is worse than silence.
    """
    notes = {}
    if not qdist:
        return notes
    hindi = (paper_language or "").lower() == "hindi"
    for qtype_frontend, info in qdist.items():
        if not isinstance(info, dict):
            continue
        attempt = info.get("attemptAny")
        if not attempt:
            continue
        sec = _normalize_qtype(qtype_frontend)
        if sec == "MCQ":
            sec = "Multiple Choice"
        printed = len(sections.get(sec) or [])
        if printed < 2 or attempt >= printed:
            continue
        notes[sec] = (
            f"{printed} में से किन्हीं {attempt} प्रश्नों के उत्तर दीजिए।"
            if hindi
            else f"Attempt any {attempt} of the {printed} questions in this section."
        )
    return notes


def _section_titles_for(paper_language: str, sections: dict | None = None) -> dict:
    """Build section headings, lettering only the sections that carry questions.

    `sections` maps section key -> list of questions. Pass it so a paper without
    MCQs starts at "Section A", not "Section B". Omit it and every section is
    lettered in the canonical order.
    """
    lang = (paper_language or "english").lower()
    names = _SECTION_NAMES.get(lang, _SECTION_NAMES["english"])
    word = _SECTION_WORD.get(lang, _SECTION_WORD["english"])

    titles = {}
    letter = 0
    for key in _SECTION_ORDER:
        if sections is not None and not sections.get(key):
            # Not rendered; keep an unlettered entry so lookups never KeyError.
            titles[key] = names[key]
            continue
        titles[key] = f"{word} {chr(65 + letter)} — {names[key]}"
        letter += 1
    return titles


def _pdf_labels_for(paper_language: str) -> dict:
    return _PDF_LABELS.get((paper_language or "english").lower(), _PDF_LABELS["english"])


def _render_paper_pdf(*, out_path, questions_sorted, school, exam_name, board,
                      class_, subject, paper_language, section_notes=None, qdist=None):
    """Render the question paper to `out_path` as PDF.

    Shared by first-time generation and by re-rendering after a manual edit, so
    an edited paper comes out looking exactly like the original.
    Pass `section_notes` to reuse stored notes; otherwise they are derived from
    `qdist`.
    """
    fonts_dir = os.path.join(current_app.root_path, "fonts")
    font_url_regular = pathlib.Path(
        os.path.join(fonts_dir, "NotoSansDevanagari-Regular.ttf")).as_uri()
    bold_path = os.path.join(fonts_dir, "NotoSansDevanagari-Bold.ttf")
    font_url_bold = (pathlib.Path(bold_path).as_uri()
                     if os.path.exists(bold_path) else font_url_regular)

    sections = {k: [] for k in _SECTION_ORDER}
    labels = _pdf_labels_for(paper_language)
    for q in questions_sorted:
        qt = _normalize_qtype(q.get("question_type", ""))
        if qt == "MCQ":
            qt = "Multiple Choice"
        if qt in sections:
            sections[qt].append(q)
    # Letter the sections only after we know which ones actually have questions.
    section_titles = _section_titles_for(paper_language, sections)
    if section_notes is None:
        section_notes = _section_choice_notes(qdist, sections, paper_language)

    env = current_app.jinja_env
    env.filters["math_render"] = render_simple_math
    rendered_html = render_template_string(
        _PAPER_HTML_TEMPLATE,
        font_url_regular=font_url_regular,
        font_url_bold=font_url_bold,
        school=school,
        exam_name=exam_name or f"{board} Board Examination",
        class_=class_, subject=subject,
        date=datetime.now().strftime("%d-%m-%Y"),
        sections=sections, section_titles=section_titles,
        section_notes=section_notes or {}, labels=labels,
        lang_attr="hi" if paper_language == "hindi" else "en",
    )
    HTML(string=rendered_html).write_pdf(out_path)


def _persist_and_render_paper(
    *, questions, board, class_, subject, school, exam_name, paper_language, visitor_id,
    qdist=None,
):
    """Persist paper to DB and render PDF/JSON/DOCX-deferred. Returns (paper_id, urls, summary)."""
    paper_id = str(uuid.uuid4())[:8]
    pdf_filename = f"paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    word_filename = f"{paper_id}.docx"
    answer_key_filename = f"answer_key_{paper_id}.pdf"

    summary = {
        "total_questions": len(questions),
        "total_marks": sum(int(q.get("marks", 0)) for q in questions),
    }

    paper_entry = Paper()
    paper_entry.paper_id = paper_id
    paper_entry.exam_name = exam_name
    paper_entry.school_name = school
    paper_entry.board = board
    paper_entry.class_ = class_
    paper_entry.subject = subject
    paper_entry.total_questions = summary["total_questions"]
    paper_entry.total_marks = summary["total_marks"]
    paper_entry.pdf_path = f"/static/papers/{pdf_filename}"
    paper_entry.word_path = f"/static/papers/{word_filename}"
    paper_entry.answer_key_path = f"/static/papers/{answer_key_filename}"
    paper_entry.visitor_id = visitor_id
    db.session.add(paper_entry)
    db.session.commit()

    for q in questions:
        pq = PaperQuestion()
        pq.paper_id = paper_id
        pq.question_id = q.get("id")
        pq.question_text = q.get("question_text") or q.get("question")
        pq.type = q.get("question_type")
        pq.difficulty = q.get("difficulty")
        pq.marks = q.get("marks")
        pq.options = q.get("options") if q.get("options") else None
        pq.answer = q.get("answer", "Not provided")
        db.session.add(pq)
    db.session.commit()

    papers_dir = os.path.join(current_app.root_path, "static", "papers")
    os.makedirs(papers_dir, exist_ok=True)
    json_path = os.path.join(papers_dir, f"{paper_id}.json")

    type_order = ["MCQ", "Multiple Choice", "Fill in the Blanks", "Fill", "Short Answer", "Short",
                  "Long Answer", "Long", "Matching", "Match", "Match the Following", "Case Study", "Case"]
    def get_type_order(q):
        t = (q.get("type") or q.get("question_type") or "").strip()
        try:
            return type_order.index(t)
        except ValueError:
            return len(type_order)

    questions_sorted = sorted(questions, key=get_type_order)
    with open(json_path, "w", encoding="utf-8") as f:
        sections_preview = {k: [] for k in _SECTION_ORDER}
        for q in questions_sorted:
            _qt = _normalize_qtype(q.get("question_type", ""))
            if _qt == "MCQ":
                _qt = "Multiple Choice"
            if _qt in sections_preview:
                sections_preview[_qt].append(q)
        json.dump({
            "paper_id": paper_id, "examName": exam_name, "schoolName": school,
            "schoolBoard": board, "class": class_, "subject": subject,
            "questions": questions_sorted, "summary": summary,
            # Persisted so the Word and answer-key exports show the same
            # "attempt any N of M" instruction as the PDF.
            "section_notes": _section_choice_notes(qdist, sections_preview, paper_language),
        }, f, indent=2, ensure_ascii=False)

    _render_paper_pdf(
        out_path=os.path.join(papers_dir, pdf_filename),
        questions_sorted=questions_sorted, school=school, exam_name=exam_name,
        board=board, class_=class_, subject=subject,
        paper_language=paper_language, section_notes=None, qdist=qdist,
    )

    urls = {
        "pdf_url": f"/static/papers/{pdf_filename}",
        "word_url": f"/api/download/word/{paper_id}",
        "answer_key_url": f"/api/download/answer_key/{paper_id}",
    }
    return paper_id, urls, summary


_PAPER_HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="{{ lang_attr }}">
<head>
<meta charset="UTF-8">
<style>
@page { size: A4; margin: 22mm 18mm; }

/* Register BOTH weights so bold runs don't fall back to a non-Devanagari face. */
@font-face {
  font-family: 'NotoDev';
  src: url('{{ font_url_regular }}') format('truetype');
  font-weight: 400;
  font-style: normal;
}
@font-face {
  font-family: 'NotoDev';
  src: url('{{ font_url_bold }}') format('truetype');
  font-weight: 700;
  font-style: normal;
}

html, body {
  font-family: 'NotoDev', 'Noto Sans Devanagari', sans-serif;
  font-size: 12pt;
  line-height: 1.55;          /* Devanagari needs extra vertical room */
  background: #fff;
  color: #1a1a1a;
}

.paper-header { text-align: center; margin-bottom: 14px; border-bottom: 2px solid #222; padding-bottom: 8px; }
.paper-header h1 { font-size: 20pt; margin: 0; font-weight: 700; }
.paper-header h2 { font-size: 14pt; margin: 4px 0; font-weight: 400; }
.paper-header h3 { font-size: 12pt; margin: 4px 0; font-weight: 400; color: #444; }

.details { display: flex; justify-content: space-between; margin-bottom: 14px; font-size: 11pt; color: #333; }

section { margin-bottom: 18px; page-break-inside: auto; }
h4.section-title {
  font-size: 12pt;
  font-weight: 700;
  border-bottom: 1px solid #bbb;
  padding: 4px 0;
  margin: 14px 0 10px;
}

ol.question-list { list-style-type: none; padding-left: 0; margin-top: 0; }

/* "Attempt any 4 of 6" — the internal-choice instruction for a section. */
p.section-note {
  margin: -4px 0 8px;
  font-size: 11pt;
  font-style: italic;
  color: #444;
}

li.question {
  margin-bottom: 12px;
  page-break-inside: avoid;
}

.question-text {
  display: flex;
  justify-content: space-between;
  align-items: flex-start;
  gap: 14px;
}
.question-text .q-body { flex: 1; }
.question-text .q-body b { font-weight: 700; margin-right: 4px; }
.question-text .marks {
  font-weight: 700;
  white-space: nowrap;
  padding-left: 12px;
  font-size: 10.5pt;
  color: #555;
}

ol.options {
  list-style-type: lower-alpha;
  padding-left: 28px;
  margin: 6px 0 0;
}
ol.options li {
  margin-bottom: 2px;
  padding-left: 4px;
}

/* Match the Following: a real two-column table, Column A numbered 1..n and
   Column B lettered a..n, so the student has something to match. */
table.match {
  width: 100%;
  border-collapse: collapse;
  margin: 8px 0 0 18px;
  font-size: 11.5pt;
}
table.match th {
  text-align: left;
  font-weight: 700;
  border-bottom: 1px solid #bbb;
  padding: 3px 8px;
  width: 50%;
}
table.match td {
  vertical-align: top;
  padding: 3px 8px;
}

/* Case Study: an indented stimulus passage, then its sub-questions. */
.case-passage {
  margin: 8px 0 6px 18px;
  padding: 8px 10px;
  border-left: 3px solid #bbb;
  background: #f7f7f7;
  text-align: justify;
}
ol.sub-questions {
  list-style-type: lower-roman;
  padding-left: 46px;
  margin: 4px 0 0;
}
ol.sub-questions li { margin-bottom: 3px; }
</style>
</head>
<body>
<div class="paper-container">
  <div class="paper-header">
    <h1>{{ school }}</h1>
    <h2>{{ exam_name }}</h2>
    <h3>{{ labels.class }} {{ class_ }} &middot; {{ subject }}</h3>
  </div>
  <div class="details">
    <span>{{ labels.date }}: {{ date }}</span>
  </div>
  {% set q_num = namespace(value=1) %}
  {% for sec_type, q_list in sections.items() %}
    {% if q_list %}
    <section>
      <h4 class="section-title">{{ section_titles[sec_type] }}</h4>
      {% if section_notes.get(sec_type) %}
      <p class="section-note">{{ section_notes[sec_type] }}</p>
      {% endif %}
      <ol class="question-list">
        {% for q in q_list %}
        <li class="question">
          <div class="question-text">
            <span class="q-body"><b>{{ labels.q_prefix }}{{ q_num.value }}.</b> {{ q.question_text | math_render | safe }}</span>
            <span class="marks">({{ q.marks }} {{ labels.marks_one if q.marks == 1 else labels.marks_many }})</span>
          </div>
          {% if q.question_type == 'MCQ' and q.options %}
          <ol class="options">
            {% for opt in q.options %}
            <li>{{ opt | math_render | safe }}</li>
            {% endfor %}
          </ol>
          {% endif %}
          {% if q.pairs %}
          <table class="match">
            <tr><th>{{ labels.column_a }}</th><th>{{ labels.column_b }}</th></tr>
            {% for pair in q.pairs %}
            <tr>
              <td>({{ loop.index }}) {{ pair[0] | math_render | safe }}</td>
              <td>({{ 'abcdefgh'[loop.index0] }}) {{ pair[1] | math_render | safe }}</td>
            </tr>
            {% endfor %}
          </table>
          {% endif %}
          {% if q.passage %}
          <div class="case-passage">{{ q.passage | math_render | safe }}</div>
          {% endif %}
          {% if q.sub_questions %}
          <ol class="sub-questions">
            {% for sq in q.sub_questions %}
            <li>{{ sq | math_render | safe }}</li>
            {% endfor %}
          </ol>
          {% endif %}
        </li>
        {% set q_num.value = q_num.value + 1 %}
        {% endfor %}
      </ol>
    </section>
    {% endif %}
  {% endfor %}
</div>
</body>
</html>
"""


def _roman(n: int) -> str:
    """Lowercase roman numeral for case-study sub-question labels (i, ii, iii...)."""
    numerals = [(10, "x"), (9, "ix"), (5, "v"), (4, "iv"), (1, "i")]
    out = ""
    for value, sym in numerals:
        while n >= value:
            out += sym
            n -= value
    return out or "i"


def _set_run_font(run, font_name: str = "Nirmala UI", size_pt: float = 11.0, bold: bool = False):
    """Apply a Devanagari-capable font to a python-docx run (both Latin and East-Asian/Indic ranges)."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), font_name)


@main.route("/api/download/word/<paper_id>", methods=["GET"])
def download_word(paper_id):
    json_path = os.path.join(current_app.root_path, "static", "papers", f"{paper_id}.json")
    if not os.path.exists(json_path):
        return jsonify({"error": "Paper not found"}), 404
    with open(json_path, "r", encoding="utf-8") as f:
        paper = json.load(f)

    is_hindi = any((q.get("language") or "").lower() == "hindi" for q in paper.get("questions", []))
    paper_language = "hindi" if is_hindi else "english"
    font_name = "Nirmala UI" if is_hindi else "Calibri"
    labels = _pdf_labels_for(paper_language)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:cs"), font_name)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    title = doc.add_paragraph()
    title_run = title.add_run(paper.get("examName") or ("प्रश्न पत्र" if is_hindi else "Question Paper"))
    _set_run_font(title_run, font_name, 18, bold=True)
    title.alignment = 1

    meta = doc.add_paragraph()
    _set_run_font(meta.add_run(f"{paper.get('schoolName', '')}"), font_name, 12, bold=True)
    meta.alignment = 1
    sub = doc.add_paragraph()
    sub_text = (f"{labels['class']} {paper.get('class', '')}  ·  {paper.get('subject', '')}  ·  "
                f"{'बोर्ड' if is_hindi else 'Board'}: {paper.get('schoolBoard', '')}")
    _set_run_font(sub.add_run(sub_text), font_name, 11)
    sub.alignment = 1

    doc.add_paragraph()

    # Group questions by section so they appear under translated section headings.
    section_order = ["MCQ", "Multiple Choice", "Fill in the Blanks",
                     "Short Answer", "Long Answer", "Matching", "Case Study"]
    section_key = {
        "MCQ": "Multiple Choice", "Multiple Choice": "Multiple Choice",
        "Fill in the Blanks": "Fill in the Blanks",
        "Short Answer": "Short Answer", "Long Answer": "Long Answer",
        "Matching": "Matching", "Case Study": "Case Study",
    }
    grouped = {k: [] for k in _SECTION_ORDER}
    for q in paper.get("questions", []):
        sk = section_key.get(_normalize_qtype(q.get("question_type", "")))
        if sk:
            grouped[sk].append(q)
    section_titles = _section_titles_for(paper_language, grouped)

    q_counter = 1
    for sec_key, q_list in grouped.items():
        if not q_list:
            continue
        h = doc.add_paragraph()
        _set_run_font(h.add_run(section_titles[sec_key]), font_name, 13, bold=True)
        note = (paper.get("section_notes") or {}).get(sec_key)
        if note:
            np_ = doc.add_paragraph()
            nr = np_.add_run(note)
            _set_run_font(nr, font_name, 10)
            nr.italic = True

        for q in q_list:
            p = doc.add_paragraph()
            _set_run_font(p.add_run(f"{labels['q_prefix']}{q_counter}. "), font_name, 11, bold=True)
            _set_run_font(p.add_run(q.get("question_text") or q.get("question") or ""), font_name, 11)
            marks = int(q.get("marks", 0) or 0)
            marks_label = labels["marks_one"] if marks == 1 else labels["marks_many"]
            _set_run_font(p.add_run(f"   ({marks} {marks_label})"), font_name, 10, bold=True)

            if _normalize_qtype(q.get("question_type", "")) == "MCQ":
                for idx, opt in enumerate(q.get("options") or [], start=1):
                    op = doc.add_paragraph()
                    _set_run_font(op.add_run(f"   ({chr(96+idx)}) "), font_name, 11, bold=True)
                    _set_run_font(op.add_run(str(opt)), font_name, 11)

            # Match the Following: a real 2-column table the student can fill in.
            pairs = q.get("pairs") or []
            if pairs:
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for cell, text in zip(hdr, (labels["column_a"], labels["column_b"])):
                    cell.paragraphs[0].runs.clear() if cell.paragraphs[0].runs else None
                    _set_run_font(cell.paragraphs[0].add_run(text), font_name, 11, bold=True)
                for idx, pair in enumerate(pairs):
                    row = table.add_row().cells
                    _set_run_font(row[0].paragraphs[0].add_run(f"({idx+1}) {pair[0]}"), font_name, 11)
                    _set_run_font(
                        row[1].paragraphs[0].add_run(f"({chr(97+idx)}) {pair[1]}"), font_name, 11)
                doc.add_paragraph()

            # Case Study: the stimulus passage, then its sub-questions.
            if q.get("passage"):
                pp = doc.add_paragraph()
                pp.paragraph_format.left_indent = Inches(0.35)
                _set_run_font(pp.add_run(str(q["passage"])), font_name, 10.5)
            for idx, sq in enumerate(q.get("sub_questions") or [], start=1):
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Inches(0.6)
                _set_run_font(sp.add_run(f"({_roman(idx)}) "), font_name, 11, bold=True)
                _set_run_font(sp.add_run(str(sq)), font_name, 11)
            q_counter += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"paper_{paper_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

_ANSWER_KEY_HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="{{ lang_attr }}">
<head>
<meta charset="UTF-8">
<style>
@page { size: A4; margin: 22mm 18mm; }
@font-face {
  font-family: 'NotoDev';
  src: url('{{ font_url_regular }}') format('truetype');
  font-weight: 400;
}
@font-face {
  font-family: 'NotoDev';
  src: url('{{ font_url_bold }}') format('truetype');
  font-weight: 700;
}
html, body {
  font-family: 'NotoDev', 'Noto Sans Devanagari', sans-serif;
  font-size: 11.5pt;
  line-height: 1.55;
  color: #1a1a1a;
}
.header { text-align: center; border-bottom: 2px solid #222; padding-bottom: 8px; margin-bottom: 14px; }
.header h1 { font-size: 18pt; margin: 0; font-weight: 700; }
.header h2 { font-size: 12pt; margin: 4px 0; font-weight: 400; color: #444; }
.meta { font-size: 10.5pt; color: #333; margin-bottom: 14px; }
.entry { margin-bottom: 14px; page-break-inside: avoid; }
.entry .q { font-weight: 700; }
.entry .a { margin-top: 4px; padding-left: 12px; }
.entry .a .label { font-weight: 700; color: #0a6e54; }
.entry .e { margin-top: 4px; padding-left: 12px; color: #444; font-size: 10.5pt; }
.entry .e .label { font-weight: 700; }
hr.sep { border: 0; border-top: 1px dashed #ccc; margin: 10px 0; }
</style>
</head>
<body>
  <div class="header">
    <h1>{{ labels.answer_key }} — {{ exam_name }}</h1>
    <h2>{{ school }}</h2>
  </div>
  <div class="meta">{{ labels.class }} {{ class_ }} &middot; {{ subject }} &middot; {{ labels.date }}: {{ date }}</div>
  {% for q in questions %}
    <div class="entry">
      <div class="q">{{ labels.q_prefix }}{{ loop.index }}. {{ q.question_text | math_render | safe }}</div>
      {% if q.question_type == 'MCQ' and q.options %}
        <ol type="a" style="margin: 4px 0 4px 24px;">
          {% for opt in q.options %}<li>{{ opt | math_render | safe }}</li>{% endfor %}
        </ol>
      {% endif %}
      {% if q.pairs %}
        <table style="width:100%; border-collapse:collapse; margin:4px 0 4px 24px;">
          <tr>
            <th style="text-align:left; padding:2px 8px; border-bottom:1px solid #bbb;">{{ labels.column_a }}</th>
            <th style="text-align:left; padding:2px 8px; border-bottom:1px solid #bbb;">{{ labels.column_b }}</th>
          </tr>
          {% for pair in q.pairs %}
          <tr>
            <td style="padding:2px 8px; vertical-align:top;">({{ loop.index }}) {{ pair[0] | math_render | safe }}</td>
            <td style="padding:2px 8px; vertical-align:top;">({{ 'abcdefgh'[loop.index0] }}) {{ pair[1] | math_render | safe }}</td>
          </tr>
          {% endfor %}
        </table>
      {% endif %}
      {% if q.passage %}
        <div style="margin:4px 0 4px 24px; padding:6px 10px; border-left:3px solid #bbb; background:#f7f7f7;">{{ q.passage | math_render | safe }}</div>
      {% endif %}
      {% if q.sub_questions %}
        <ol type="i" style="margin:4px 0 4px 44px;">
          {% for sq in q.sub_questions %}<li>{{ sq | math_render | safe }}</li>{% endfor %}
        </ol>
      {% endif %}
      <div class="a"><span class="label">{{ labels.answer }}:</span> {{ q.answer | math_render | safe }}</div>
      {% if q.explanation %}
        <div class="e"><span class="label">{{ labels.explanation }}:</span> {{ q.explanation | math_render | safe }}</div>
      {% endif %}
    </div>
    <hr class="sep">
  {% endfor %}
</body>
</html>
"""


@main.route("/api/paper/<paper_id>", methods=["PATCH"])
@limiter.limit("60 per hour")
def update_paper(paper_id):
    """Save manual edits to a generated paper and re-render its PDF.

    Marks the paper as edited, which permanently withdraws its answer key: the
    stored answers belong to the questions the AI wrote, and once a teacher
    rewrites a question we can no longer claim to know the answer.
    """
    papers_dir = os.path.join(current_app.root_path, "static", "papers")
    json_path = os.path.join(papers_dir, f"{paper_id}.json")
    if not os.path.exists(json_path):
        return jsonify({"error": "Paper not found"}), 404

    payload = request.get_json(silent=True) or {}
    edited = payload.get("questions")
    if not isinstance(edited, list) or not edited:
        return jsonify({"error": "No questions supplied"}), 400
    if len(edited) > 100:
        return jsonify({"error": "Too many questions"}), 400

    with open(json_path, "r", encoding="utf-8") as f:
        paper = json.load(f)

    original = paper.get("questions", [])
    merged = []
    for item in edited:
        if not isinstance(item, dict):
            continue
        idx = item.get("index")
        base = dict(original[idx]) if isinstance(idx, int) and 0 <= idx < len(original) else {}

        text = str(item.get("question_text", base.get("question_text", "")))[:4000].strip()
        if not text:
            continue                       # a question emptied out is a deletion
        base["question_text"] = text
        base["question"] = text
        try:
            base["marks"] = max(0, min(20, int(item.get("marks", base.get("marks", 1)))))
        except (TypeError, ValueError):
            base["marks"] = base.get("marks", 1)

        if isinstance(item.get("options"), list):
            opts = [str(o).strip()[:500] for o in item["options"] if str(o).strip()]
            base["options"] = opts[:6] or None
        if isinstance(item.get("pairs"), list):
            pairs = []
            for pr in item["pairs"]:
                if isinstance(pr, (list, tuple)) and len(pr) >= 2:
                    left, right = str(pr[0]).strip()[:300], str(pr[1]).strip()[:300]
                    if left or right:
                        pairs.append([left, right])
            base["pairs"] = pairs[:6] or None
        if isinstance(item.get("passage"), str):
            base["passage"] = item["passage"].strip()[:4000] or None
        if isinstance(item.get("sub_questions"), list):
            subs = [str(x).strip()[:500] for x in item["sub_questions"] if str(x).strip()]
            base["sub_questions"] = subs[:6] or None

        base["source"] = "Edited"
        merged.append(base)

    if not merged:
        return jsonify({"error": "Every question was empty"}), 400

    paper["questions"] = merged
    paper["edited"] = True
    paper["summary"] = {
        "total_questions": len(merged),
        "total_marks": sum(int(q.get("marks", 0) or 0) for q in merged),
    }

    paper_language = "english"
    for q in merged:
        if (q.get("language") or "").lower() == "hindi":
            paper_language = "hindi"
            break

    # New filename each save so browsers and CDNs cannot serve a stale PDF.
    pdf_filename = f"paper_{paper_id}_{datetime.now().strftime('%H%M%S')}.pdf"
    try:
        _render_paper_pdf(
            out_path=os.path.join(papers_dir, pdf_filename),
            questions_sorted=merged,
            school=paper.get("schoolName", ""), exam_name=paper.get("examName", ""),
            board=paper.get("schoolBoard", ""), class_=paper.get("class", ""),
            subject=paper.get("subject", ""), paper_language=paper_language,
            section_notes=paper.get("section_notes") or {},
        )
    except Exception as e:
        current_app.logger.error("Re-render after edit failed for %s: %s", paper_id, e)
        return jsonify({"error": "Could not re-render the paper."}), 500

    paper["pdf_url"] = f"/static/papers/{pdf_filename}"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(paper, f, indent=2, ensure_ascii=False)

    row = Paper.query.filter_by(paper_id=paper_id).first()
    if row:
        row.pdf_path = paper["pdf_url"]
        row.total_questions = paper["summary"]["total_questions"]
        row.total_marks = paper["summary"]["total_marks"]
        PaperQuestion.query.filter_by(paper_id=paper_id).delete()
        for q in merged:
            pq = PaperQuestion()
            pq.paper_id = paper_id
            pq.question_id = q.get("id")
            pq.question_text = q.get("question_text")
            pq.type = q.get("question_type")
            pq.difficulty = q.get("difficulty")
            pq.marks = q.get("marks")
            pq.options = q.get("options") or None
            pq.answer = q.get("answer", "Not provided")
            db.session.add(pq)
        db.session.commit()

    return jsonify({
        "paper_id": paper_id,
        "edited": True,
        "summary": paper["summary"],
        "pdf_url": paper["pdf_url"],
        "word_url": f"/api/download/word/{paper_id}",
        "answer_key_available": False,
    })


@main.route("/api/download/answer_key/<paper_id>", methods=["GET"])
def download_answer_key(paper_id):
    json_path = os.path.join(current_app.root_path, "static", "papers", f"{paper_id}.json")
    if not os.path.exists(json_path):
        return jsonify({"error": "Paper not found"}), 404

    with open(json_path, "r", encoding="utf-8") as f:
        paper = json.load(f)

    # The button is hidden client-side after an edit; this is the matching
    # server-side guard so the URL cannot simply be requested directly.
    if paper.get("edited"):
        return jsonify({
            "error": "This paper was edited manually, so its answer key is no "
                     "longer available.",
        }), 409

    # Detect paper language from any question entry (defaults to English).
    paper_language = "english"
    for q in paper.get("questions", []):
        if (q.get("language") or "").lower() == "hindi":
            paper_language = "hindi"
            break
    lang = "hi" if paper_language == "hindi" else "en"

    fonts_dir = os.path.join(current_app.root_path, "fonts")
    font_url_regular = pathlib.Path(os.path.join(fonts_dir, "NotoSansDevanagari-Regular.ttf")).as_uri()
    bold_path = os.path.join(fonts_dir, "NotoSansDevanagari-Bold.ttf")
    font_url_bold = pathlib.Path(bold_path).as_uri() if os.path.exists(bold_path) else font_url_regular

    env = current_app.jinja_env
    env.filters["math_render"] = render_simple_math

    rendered_html = render_template_string(
        _ANSWER_KEY_HTML_TEMPLATE,
        questions=paper.get("questions", []),
        exam_name=paper.get("examName") or ("परीक्षा" if paper_language == "hindi" else "Exam"),
        school=paper.get("schoolName") or "",
        class_=paper.get("class") or "",
        subject=paper.get("subject") or "",
        date=datetime.now().strftime("%d-%m-%Y"),
        font_url_regular=font_url_regular,
        font_url_bold=font_url_bold,
        lang_attr=lang,
        labels=_pdf_labels_for(paper_language),
    )

    pdf_bytes = HTML(string=rendered_html).write_pdf()
    buf = io.BytesIO(pdf_bytes)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"answer_key_{paper_id}.pdf",
        mimetype="application/pdf",
    )



